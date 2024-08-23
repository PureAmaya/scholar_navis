import os
import random
from docx import Document
from .tools.article_library_ctrl import SQLiteDatabase,db_type,get_tmp_dir_of_this_user, command_information, check_library_exist_and_assistant, promote_file_to_downloadzone, VERSION,pdf_reader,read_and_clean_pdf_text,breakdown_text_to_satisfy_token_limit,article_library_ctrl, CatchException,get_files_from_everything,get_log_folder,get_user,l,update_ui,update_ui_lastest_msg,request_gpt_model_in_new_thread_with_ui_alive,GPT_prefer_language

# ! 后面再说，这个目前没啥必要
# ! 没改好，也暂时没啥必要

@check_library_exist_and_assistant(accept_nonexistent=True,accept_blank=True)
@CatchException
def 更好的文章润色(txt: str, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    """
        通过用户上传文章（1篇），让AI阅读这些文章，并模仿该文章的语言风格、词语用法、叙事逻辑，时态语态，进行润色
        就1篇吧，毕竟还得上传一份用户的手稿，别超过token了
        
    """
    feature = command_information.更好的文章润色_help

    chatbot.append([feature, l("请将参考的文献和草稿压缩为压缩包上传")])
    
    chatbot.append([l('参考文章要求：pdf格式，是英文，最好是该领域内的、由英语母语作者撰写的文章'), 
                    l("草稿要求：非pdf格式（docx、txt均可），英文最佳，其他语言亦可。无需遵守论文的格式，仅仅是对于语言进行润色")])
    yield from update_ui(chatbot=chatbot, history=[])
    
    if txt == '':return
    elif not os.path.exists(txt):
        chatbot.append([txt, l("上传文件不存在，请检查路径或重新上传")])
        yield from update_ui(chatbot=chatbot, history=[])
        return
    
    _,ref_pdf,_ = get_files_from_everything(txt,'.pdf')
    chatbot.append([l('已接收到上传的文件'),'...'])
    yield from update_ui(chatbot=chatbot, history=[])
    # 如果上传了多个pdf，提醒一下，我们只会要第一个
    only_one_ref_indicator = l('上传的参考文章多于1篇，只会使用第一篇文章\n') if len(ref_pdf) > 2 else ''
    # 参考文章读取
    ref_pdf_content ,_ = read_and_clean_pdf_text(ref_pdf[0]) 
    # 参考文章的doi
    doi = pdf_reader.get_pdf_inf(ref_pdf[0])
    # 草稿
    upload_dir = os.path.dirname(ref_pdf[0])
    draft_content = ''
    
    for file in os.listdir(upload_dir):
        if file.endswith('doc'):
            yield from update_ui_lastest_msg(l('doc格式过于老旧，请将其转换为docx再使用'),chatbot,[])
            return
        elif file.endswith('docx'):
            document = Document(os.path.join(upload_dir,file))
            for paragraph in document.paragraphs:draft_content = f'{draft_content}\n{paragraph.text}'
        elif file.endswith('txt'):
            with open(os.path.join(upload_dir,file),'r') as f:
                draft_content = f.read()
    
    if len(ref_pdf_content) < 1000:
        yield from update_ui_lastest_msg(l('上传的参考文本长度不足或读取失败，请更换参考文献'),chatbot,[])
        return
    
    chatbot.clear()
    chatbot.append([l('{}参考文献和草稿获取成功，正在执行润色操作...').format(only_one_ref_indicator), 
                    l("请等待...")])
    yield from update_ui(chatbot=chatbot, history=[])
    
    # 后面请求LLM用的特定提示词
    prompt = "You're a native English speaker. Just give me the results in American English straight up. Keep everything formal; no slang or casual language."
    
    # 把参考文献切割，防止超出token
    ref_pdf_content_broken = breakdown_text_to_satisfy_token_limit(txt=ref_pdf_content, limit=2048, llm_model=llm_kwargs['llm_model'])
    # 把手稿切割，便于不同的切割内容进行润色对照
    draft_content_broken = breakdown_text_to_satisfy_token_limit(txt=draft_content,limit=2048,llm_model=llm_kwargs['llm_model'])

    # 记录LLM学习的语言模板
    emulated_reference = ''
    
    # 看看有没有之前的储存，有的话就直接用了，不用请求LLM了
    with SQLiteDatabase(type=db_type.doi_emulate_polish) as ft:
        emulated_reference = ft.select(doi,('emulated_content',))

    # 没有储存的话，就请求吧
    if emulated_reference == '' or emulated_reference is None:
        for index, content in enumerate(ref_pdf_content_broken):
            i_say = f"please answer me in {GPT_prefer_language}.\
                    Go ahead and read the passage, then sum it up, keeping the same style, word choices, story flow, and tenses and voice: {content}"
            i_say_show = l("[{i}/{total}] AI阅读中...").format(i=index+1,total=len(ref_pdf_content_broken))
            
            #通常的，对分片进行总结
            gpt_say_for_fragment = yield from request_gpt_model_in_new_thread_with_ui_alive(i_say, i_say_show, llm_kwargs, chatbot,
                                                                            history=[], # 分段总结的时候，不需要历史记录
                                                                            sys_prompt=prompt)
            history.append(gpt_say_for_fragment)
        
        chatbot.append([l('参考文章学习完成'),l('下面进行润色')])
        yield from update_ui(chatbot=chatbot, history=[])
        
        # 记录LLM学习的语言模板
        emulated_reference = '\n\n\n'.join(history)
        
        # 看完了参考文章，先记录一下，尽可能不请求LLM
        with SQLiteDatabase(type=db_type.doi_emulate_polish) as ft:
            ft.insert_ingore(doi,('emulated_content',),(emulated_reference,))
    
    # 润色输出内容
    polish_export = []
    
    #print(f'ref: {ref_pdf_content_broken[0]}')
    
    # 对手稿的分片进行润色
    for index, content in enumerate(draft_content_broken):
        
        i_say = f"Make sure to spruce up this text. If it's not in English, translate it into English while you're at it and make it look real nice： {content}"
        i_say_show = l("[{i}/{total}] AI润色中...").format(i=index+1,total=len(draft_content_broken))
        
        #通常的，对分片进行总结
        gpt_say_for_fragment = yield from request_gpt_model_in_new_thread_with_ui_alive(i_say, i_say_show, llm_kwargs, chatbot,
                                                                        # history 用于提醒模拟
                                                                        history=[f"When you are spiffing up the text, you gotta copy the style, the way words are used, the story's flow, and the tenses and voices from the stuff below: {emulated_reference}"], # 提醒模仿
                                                                        sys_prompt=prompt)
        polish_export.extend([content,gpt_say_for_fragment])
        
    warning_msg = l('请注意，AI可能会有错误，务必人工校准！为防止直接套用，输出的内容中已经添加无用文本')
    chatbot.append([l('润色完成'),l('下面是润色结果')+f'<br><font color=red><b>{warning_msg}</b></font>'])
    yield from update_ui(chatbot=chatbot, history=[])

    # 整理输出
    docx_output_fp = os.path.join(get_tmp_dir_of_this_user(chatbot,'polish',[]),'polished.docx')
    document = Document()
    # 第一页，放信息
    document.add_heading(l('通过 Scholar Navis 进行AI润色'), 0)
    document.add_paragraph(f'Scholar Navis ver.{VERSION}')
    document.add_paragraph('')
    document.add_paragraph(feature) # 复述一下功能
    document.add_paragraph(warning_msg) # 复述一下警告
    document.add_paragraph(l('下一页为润色内容。上方的文本为原文（加粗），下方的文本是润色后的内容'))
    document.add_page_break()
    for index, content in enumerate(polish_export):
        # 原始文本
        if index % 2 == 0:
            document.add_paragraph(content)
            document.add_paragraph('')
        # 润色内容
        else:
            # 先找到第一个空格
            blank_index = content.find(' ')
            # 储存所有的空格位置
            spaces_positions = []
            # 循环查找所有的空格位置
            while blank_index != -1:
                spaces_positions.append(blank_index)
                blank_index = content.find(' ',blank_index + 1)
                
            watermark = f' AI-GENERATED' # 防君子不防小人。即使有一个完美的加水印方法，也可以靠改代码去除
            # 替换哪个空格
            index = random.randint(0,len(spaces_positions) - 1)
            # 随机把空格替换成水印
            content = content[:spaces_positions[index]] + watermark + content[spaces_positions[index]:] 
            document.add_paragraph(content)
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
    # 保存
    document.save(docx_output_fp)
    
    promote_file_to_downloadzone(file=docx_output_fp,chatbot=chatbot)
